"""
Document Parser for extracting questions and answers from .docx and .pdf files
"""
import re
import json
import logging
from typing import List, Dict, Tuple, Optional, TYPE_CHECKING
from dataclasses import dataclass
from pathlib import Path

# Document processing libraries
import docx
import pdfplumber
from docx import Document
from docx.shared import Inches

# Django imports
from django.core.files.base import ContentFile
from django.utils import timezone

if TYPE_CHECKING:
    from .models import Assessment

logger = logging.getLogger(__name__)

@dataclass
class ExtractedQuestion:
    """Data class for extracted questions"""
    question_text: str
    question_type: str
    points: int = 1
    difficulty: str = 'medium'
    choices: List[Dict] = None
    correct_answers: List[str] = None
    explanation: str = ''
    hint: str = ''
    order: int = 0
    confidence: float = 0.0
    
    def __post_init__(self):
        if self.choices is None:
            self.choices = []
        if self.correct_answers is None:
            self.correct_answers = []

@dataclass
class ParseResult:
    """Result of document parsing"""
    questions: List[ExtractedQuestion]
    total_questions: int
    questions_with_answers: int
    processing_log: Dict
    errors: List[str]
    confidence_score: float

class QuestionPatterns:
    """Regular expression patterns for detecting questions and answers"""
    
    # Question number patterns
    QUESTION_NUMBERS = [
        r'^\s*(\d+)\.\s*',  # 1. Question
        r'^\s*(\d+)\)\s*',  # 1) Question
        r'^\s*Question\s+(\d+)[:.]?\s*',  # Question 1: or Question 1.
        r'^\s*Q\.?\s*(\d+)[:.]?\s*',  # Q.1: or Q1.
        r'^\s*No\.?\s*(\d+)[:.]?\s*',  # No.1: or No 1.
    ]
    
    # Multiple choice patterns
    CHOICE_PATTERNS = [
        r'^\s*([A-Z])\.\s*(.+)$',  # A. Choice
        r'^\s*([A-Z])\)\s*(.+)$',  # A) Choice
        r'^\s*\(([A-Z])\)\s*(.+)$',  # (A) Choice
        r'^\s*([a-z])\.\s*(.+)$',  # a. Choice
        r'^\s*([a-z])\)\s*(.+)$',  # a) Choice
        r'^\s*\(([a-z])\)\s*(.+)$',  # (a) Choice
    ]
    
    # Answer key patterns
    ANSWER_PATTERNS = [
        r'^\s*Answer[s]?[:.]?\s*([A-Za-z0-9,\s]+)$',  # Answer: A, B
        r'^\s*Correct[:\s]+([A-Za-z0-9,\s]+)$',  # Correct: A
        r'^\s*Key[:.]?\s*([A-Za-z0-9,\s]+)$',  # Key: A
        r'^\s*(\d+)\.\s*([A-Za-z])\s*$',  # 1. A (answer key format)
    ]
    
    # True/False patterns
    TRUE_FALSE_PATTERNS = [
        r'\b(True|False)\b',
        r'\b(T|F)\b',
        r'\b(Correct|Incorrect)\b',
        r'\b(Yes|No)\b',
    ]
    
    # Fill in the blank patterns
    BLANK_PATTERNS = [
        r'_{3,}',  # Three or more underscores
        r'\[.*?\]',  # [blank]
        r'\(.*?\)',  # (blank)
    ]
    
    # Question type indicators
    TYPE_INDICATORS = {
        'multiple_choice': [
            r'\bchoose\b', r'\bselect\b', r'\bmultiple choice\b',
            r'\b[A-D]\)', r'\b[A-D]\.', r'\boptions?\b'
        ],
        'true_false': [
            r'\btrue\s+or\s+false\b', r'\bt/f\b', r'\btrue\s*[/\\]\s*false\b'
        ],
        'identification': [
            r'\bidentify\b', r'\bname\b', r'\bwhat\s+is\b', r'\bdefine\b'
        ],
        'enumeration': [
            r'\benumerate\b', r'\blist\b', r'\bgive\s+\d+\b', r'\bname\s+\d+\b'
        ],
        'essay': [
            r'\bexplain\b', r'\bdiscuss\b', r'\bdescribe\b', r'\banalyze\b',
            r'\bevaluate\b', r'\bcompare\b', r'\bcontrast\b'
        ]
    }

class DocumentParser:
    """Base document parser class"""
    
    def __init__(self):
        self.patterns = QuestionPatterns()
        self.processing_log = {
            'start_time': timezone.now().isoformat(),
            'steps': [],
            'warnings': [],
            'statistics': {}
        }
        self.errors = []
    
    def log_step(self, step: str, details: str = ''):
        """Log a processing step"""
        self.processing_log['steps'].append({
            'step': step,
            'details': details,
            'timestamp': timezone.now().isoformat()
        })
        logger.info(f"Parser step: {step} - {details}")
    
    def log_warning(self, warning: str):
        """Log a warning"""
        self.processing_log['warnings'].append(warning)
        logger.warning(f"Parser warning: {warning}")
    
    def log_error(self, error: str):
        """Log an error"""
        self.errors.append(error)
        logger.error(f"Parser error: {error}")
    
    def detect_question_type(self, text: str) -> str:
        """Detect question type based on text content"""
        text_lower = text.lower()
        
        # Check for specific patterns
        for q_type, indicators in self.patterns.TYPE_INDICATORS.items():
            for indicator in indicators:
                if re.search(indicator, text_lower):
                    return q_type
        
        # Check for true/false patterns
        if any(re.search(pattern, text, re.IGNORECASE) 
               for pattern in self.patterns.TRUE_FALSE_PATTERNS):
            return 'true_false'
        
        # Check for blanks
        if any(re.search(pattern, text) 
               for pattern in self.patterns.BLANK_PATTERNS):
            return 'fill_blank'
        
        # Default to multiple choice if choices are found, otherwise identification
        return 'multiple_choice'
    
    def extract_choices(self, lines: List[str], start_idx: int) -> Tuple[List[Dict], int]:
        """Extract multiple choice options from lines"""
        choices = []
        current_idx = start_idx
        
        for i in range(start_idx, len(lines)):
            line = lines[i].strip()
            if not line:
                continue
                
            # Check if this line matches a choice pattern
            choice_match = None
            for pattern in self.patterns.CHOICE_PATTERNS:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    choice_match = match
                    break
            
            if choice_match:
                choice_letter = choice_match.group(1).upper()
                choice_text = choice_match.group(2).strip()
                choices.append({
                    'letter': choice_letter,
                    'text': choice_text,
                    'is_correct': False  # Will be determined by answer key
                })
                current_idx = i + 1
            else:
                # Stop if we hit a line that doesn't look like a choice
                if choices:  # Only stop if we've found at least one choice
                    break
                current_idx = i + 1
        
        return choices, current_idx
    
    def parse_answer_key(self, text: str) -> Dict[int, List[str]]:
        """Parse answer key from text"""
        answer_key = {}
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try different answer patterns
            for pattern in self.patterns.ANSWER_PATTERNS:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    # Extract question number and answer
                    if len(match.groups()) >= 2:
                        q_num = int(match.group(1))
                        answer = match.group(2).strip()
                    else:
                        # Pattern like "Answer: A" - need to associate with question
                        answer = match.group(1).strip()
                        # Use line number or context to determine question
                        continue
                    
                    # Parse multiple answers (e.g., "A, B, C")
                    answers = [ans.strip().upper() for ans in answer.split(',')]
                    answer_key[q_num] = answers
        
        return answer_key
    
    def calculate_confidence(self, questions: List[ExtractedQuestion]) -> float:
        """Calculate overall confidence score for parsing"""
        if not questions:
            return 0.0
        
        total_confidence = 0.0
        factors = []
        
        for question in questions:
            q_confidence = 0.5  # Base confidence
            
            # Increase confidence if question has proper structure
            if question.question_text and len(question.question_text) > 10:
                q_confidence += 0.2
            
            # Increase confidence for multiple choice with choices
            if question.question_type == 'multiple_choice' and question.choices:
                q_confidence += 0.2
                
                # More confidence if choices are properly formatted
                if all('letter' in choice and 'text' in choice 
                       for choice in question.choices):
                    q_confidence += 0.1
            
            # Increase confidence if answers are provided
            if question.correct_answers:
                q_confidence += 0.2
            
            question.confidence = min(q_confidence, 1.0)
            total_confidence += question.confidence
            factors.append(question.confidence)
        
        return total_confidence / len(questions) if questions else 0.0

class DocxParser(DocumentParser):
    """Parser for .docx files"""
    
    def parse(self, file_path: str) -> ParseResult:
        """Parse a .docx file and extract questions"""
        try:
            self.log_step("Starting DOCX parsing", f"File: {file_path}")
            
            # Open the document
            doc = Document(file_path)
            
            # Extract all text content
            all_text = []
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            
            self.log_step("Text extraction complete", f"Found {len(all_text)} text elements")
            
            # Parse questions from text
            questions = self._parse_questions_from_text(all_text)
            
            # Calculate final confidence
            confidence = self.calculate_confidence(questions)
            
            self.processing_log['end_time'] = timezone.now().isoformat()
            self.processing_log['statistics'] = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'questions_found': len(questions),
                'confidence_score': confidence
            }
            
            return ParseResult(
                questions=questions,
                total_questions=len(questions),
                questions_with_answers=len([q for q in questions if q.correct_answers]),
                processing_log=self.processing_log,
                errors=self.errors,
                confidence_score=confidence
            )
            
        except Exception as e:
            self.log_error(f"Failed to parse DOCX file: {str(e)}")
            return ParseResult(
                questions=[],
                total_questions=0,
                questions_with_answers=0,
                processing_log=self.processing_log,
                errors=self.errors,
                confidence_score=0.0
            )
    
    def _parse_questions_from_text(self, text_lines: List[str]) -> List[ExtractedQuestion]:
        """Parse questions from extracted text lines"""
        questions = []
        current_question = None
        question_order = 1
        i = 0
        
        while i < len(text_lines):
            line = text_lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Check if this line starts a new question
            question_match = None
            for pattern in self.patterns.QUESTION_NUMBERS:
                match = re.match(pattern, line)
                if match:
                    question_match = match
                    break
            
            if question_match:
                # Save previous question if exists
                if current_question:
                    questions.append(current_question)
                
                # Start new question
                question_num = int(question_match.group(1))
                question_text = re.sub(r'^\s*\d+[.)]\s*', '', line).strip()
                
                # Detect question type
                question_type = self.detect_question_type(question_text)
                
                current_question = ExtractedQuestion(
                    question_text=question_text,
                    question_type=question_type,
                    order=question_order,
                    points=1
                )
                question_order += 1
                
                # Extract choices if it's a multiple choice question
                if question_type == 'multiple_choice':
                    choices, next_i = self.extract_choices(text_lines, i + 1)
                    current_question.choices = choices
                    i = next_i
                else:
                    i += 1
            else:
                # This line might be a continuation or answer
                if current_question:
                    # Check if it's an answer pattern
                    for pattern in self.patterns.ANSWER_PATTERNS:
                        match = re.match(pattern, line, re.IGNORECASE)
                        if match:
                            answer_text = match.group(1).strip() if len(match.groups()) >= 1 else ''
                            if answer_text:
                                current_question.correct_answers = [answer_text.upper()]
                            break
                    else:
                        # Append to question text if it doesn't look like an answer
                        if len(current_question.question_text) < 500:  # Reasonable limit
                            current_question.question_text += ' ' + line
                
                i += 1
        
        # Don't forget the last question
        if current_question:
            questions.append(current_question)
        
        # Post-process: match answers to choices for multiple choice questions
        self._match_answers_to_choices(questions)
        
        return questions
    
    def _match_answers_to_choices(self, questions: List[ExtractedQuestion]):
        """Match correct answers to multiple choice options"""
        for question in questions:
            if question.question_type == 'multiple_choice' and question.correct_answers and question.choices:
                correct_letters = [ans.upper() for ans in question.correct_answers]
                
                for choice in question.choices:
                    if choice.get('letter', '').upper() in correct_letters:
                        choice['is_correct'] = True

class PdfParser(DocumentParser):
    """Parser for .pdf files"""
    
    def parse(self, file_path: str) -> ParseResult:
        """Parse a .pdf file and extract questions"""
        try:
            self.log_step("Starting PDF parsing", f"File: {file_path}")
            
            all_text = []
            page_count = 0
            
            # Extract text using pdfplumber
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Split into lines and add to all_text
                            lines = page_text.split('\n')
                            all_text.extend(lines)
                        
                        self.log_step(f"Processed page {page_num + 1}", f"Extracted {len(lines) if page_text else 0} lines")
                        
                    except Exception as e:
                        self.log_warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
            
            self.log_step("Text extraction complete", f"Processed {page_count} pages, {len(all_text)} text lines")
            
            # Parse questions from text
            questions = self._parse_questions_from_text(all_text)
            
            # Calculate final confidence
            confidence = self.calculate_confidence(questions)
            
            self.processing_log['end_time'] = timezone.now().isoformat()
            self.processing_log['statistics'] = {
                'total_pages': page_count,
                'total_lines': len(all_text),
                'questions_found': len(questions),
                'confidence_score': confidence
            }
            
            return ParseResult(
                questions=questions,
                total_questions=len(questions),
                questions_with_answers=len([q for q in questions if q.correct_answers]),
                processing_log=self.processing_log,
                errors=self.errors,
                confidence_score=confidence
            )
            
        except Exception as e:
            self.log_error(f"Failed to parse PDF file: {str(e)}")
            return ParseResult(
                questions=[],
                total_questions=0,
                questions_with_answers=0,
                processing_log=self.processing_log,
                errors=self.errors,
                confidence_score=0.0
            )
    
    def _parse_questions_from_text(self, text_lines: List[str]) -> List[ExtractedQuestion]:
        """Parse questions from extracted text lines (same as DOCX for now)"""
        # Reuse the same logic as DOCX parser
        docx_parser = DocxParser()
        return docx_parser._parse_questions_from_text(text_lines)

class DocumentImportService:
    """Main service for importing documents and creating assessments"""
    
    def __init__(self):
        self.parsers = {
            'docx': DocxParser(),
            'pdf': PdfParser()
        }
    
    def process_document(self, document_import_obj) -> ParseResult:
        """Process a document import object"""
        from .models import DocumentImport  # Avoid circular import
        
        file_path = document_import_obj.document.path
        file_extension = Path(file_path).suffix.lower().lstrip('.')
        
        if file_extension not in self.parsers:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Update import status
        document_import_obj.status = 'processing'
        document_import_obj.save()
        
        try:
            # Parse the document
            parser = self.parsers[file_extension]
            result = parser.parse(file_path)
            
            # Update import object with results
            document_import_obj.questions_extracted = result.total_questions
            document_import_obj.questions_with_answers = result.questions_with_answers
            document_import_obj.processing_log = result.processing_log
            document_import_obj.processed_at = timezone.now()
            
            if result.errors:
                document_import_obj.error_messages = '\n'.join(result.errors)
                document_import_obj.status = 'needs_review' if result.questions else 'failed'
            else:
                document_import_obj.status = 'completed'
            
            document_import_obj.save()
            
            return result
            
        except Exception as e:
            document_import_obj.status = 'failed'
            document_import_obj.error_messages = str(e)
            document_import_obj.processed_at = timezone.now()
            document_import_obj.save()
            raise
    
    def create_assessment_from_import(self, document_import_obj, assessment_data: Dict) -> 'Assessment':
        """Create an assessment from parsed document data"""
        from .models import Assessment, Question, Choice, CorrectAnswer  # Avoid circular import
        
        # Process the document first if not already done
        if document_import_obj.status == 'pending':
            result = self.process_document(document_import_obj)
        else:
            # Re-process to get the questions
            result = self.process_document(document_import_obj)
        
        if not result.questions:
            raise ValueError("No questions found in the document")
        
        # Create the assessment
        assessment = Assessment.objects.create(
            title=assessment_data.get('title', f"Imported from {document_import_obj.document.name}"),
            description=assessment_data.get('description', ''),
            assessment_type=assessment_data.get('assessment_type', 'quiz'),
            creator=document_import_obj.uploaded_by,
            created_from_import=document_import_obj,
            status='draft'
        )
        
        # Create questions
        for extracted_q in result.questions:
            question = Question.objects.create(
                assessment=assessment,
                question_type=extracted_q.question_type,
                question_text=extracted_q.question_text,
                points=extracted_q.points,
                order=extracted_q.order,
                explanation=extracted_q.explanation,
                hint=extracted_q.hint,
                difficulty_level=extracted_q.difficulty,
                imported_from_document=True,
                import_confidence=extracted_q.confidence
            )
            
            # Create choices for multiple choice questions
            if extracted_q.question_type == 'multiple_choice' and extracted_q.choices:
                for i, choice_data in enumerate(extracted_q.choices):
                    Choice.objects.create(
                        question=question,
                        choice_text=choice_data.get('text', ''),
                        is_correct=choice_data.get('is_correct', False),
                        order=i + 1
                    )
            
            # Create correct answers for other question types
            if extracted_q.correct_answers and extracted_q.question_type != 'multiple_choice':
                for i, answer_text in enumerate(extracted_q.correct_answers):
                    CorrectAnswer.objects.create(
                        question=question,
                        answer_text=answer_text,
                        order=i + 1
                    )
        
        # Update document import with created assessment
        document_import_obj.created_assessment = assessment
        document_import_obj.save()
        
        # Calculate total points
        assessment.calculate_total_points()
        
        return assessment